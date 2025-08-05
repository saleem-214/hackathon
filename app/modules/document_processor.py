import fitz # PyMuPDF
import docx
import extract_msg
import os
import requests # New library for handling URLs
from io import BytesIO

def extract_text_from_pdf(file_path):
    text = ""
    if isinstance(file_path, str):
        doc = fitz.open(file_path)
    elif isinstance(file_path, BytesIO):
            doc = fitz.open(stream=file_path.read(), filetype="pdf")
    else:
        raise ValueError("Invalid input for PDF extraction")

    with doc:
        for page in doc:
            text += page.get_text()
    return text

def extract_text_from_docx(file_path):
    # Also handle byte streams for DOCX
    if isinstance(file_path, str):
        doc = docx.Document(file_path)
    elif isinstance(file_path, BytesIO):
        doc = docx.Document(file_path)
    else:
        raise ValueError("Invalid input for DOCX extraction")

    return "\n".join([para.text for para in doc.paragraphs])

def extract_text_from_msg(file_path):
    # MSG files are typically local, but this function remains
    msg = extract_msg.Message(file_path)
    return f"Subject: {msg.subject}\n\n{msg.body}"

def extract_text_from_url(url):
    print(f"Downloading document from URL: {url}")
    response = requests.get(url)
    response.raise_for_status() # Raise an exception for bad status codes
    
    file_stream = BytesIO(response.content)
    
    # Determine file type from URL or content-type header
    content_type = response.headers.get('Content-Type', '')
    if 'pdf' in url.lower() or 'pdf' in content_type.lower():
        return extract_text_from_pdf(file_stream)
    elif 'docx' in url.lower() or 'word' in content_type.lower():
        return extract_text_from_docx(file_stream)
    else:
        raise ValueError("Unsupported file format from URL: " + url)

def extract_text(source):
    if source.startswith('http'):
        return extract_text_from_url(source)
    else:
        ext = os.path.splitext(source)[-1].lower()
        if ext == ".pdf":
            return extract_text_from_pdf(source)
        elif ext == ".docx":
            return extract_text_from_docx(source)
        elif ext == ".msg":
            return extract_text_from_msg(source)
        else:
            raise ValueError("Unsupported file format: " + ext)